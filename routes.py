import json
import uuid
import io
from datetime import datetime
from schemas import RunRequest, FeedbackRequest, PromptVersionRequest, RefinementRequest, AuditUpdateRequest, ClarifyRequest, ClarifyAnswerRequest


from fastapi import APIRouter, UploadFile, File, HTTPException
from service import (
    orchestrator,
    get_db,
    policy_collection,
    AI_TOOLS_REGISTRY,
    SYSTEM_VERSION,
    call_llm,          # reuse existing LLM helper
)

router = APIRouter()


# ══════════════════════════════════════════════════════════════════════════════
# FILE TEXT EXTRACTION HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        parts  = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(parts)
    except ImportError:
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
        except ImportError:
            raise HTTPException(
                500,
                "PDF parsing library not installed. Run: pip install pypdf"
            )


def _extract_docx_text(content: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
        doc   = Document(io.BytesIO(content))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        raise HTTPException(
            500,
            "python-docx library not installed. Run: pip install python-docx"
        )


# ══════════════════════════════════════════════════════════════════════════════
# ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════════
@router.post("/api/run")
async def run_orchestrator(req: RunRequest):
    if not req.user_input.strip():
        raise HTTPException(400, "Input cannot be empty")

    result = orchestrator.invoke({
        "user_input":       req.user_input,
        "role":             req.role or "general",
        "task_type":        req.task_type or "general",
        "data_sensitivity": req.data_sensitivity or "general",
        "intent": "", "industry": "",
        "recommended_tool": "", "tool_reason": "",
        "tool_confidence": "", "tool_confidence_pct": 0,
        "tool_confidence_explanation": "",
        "tool_alternatives": [], "tool_alternative_reasons": [],
        "tool_alternative_confidence_pcts": [],
        "policy_flags": [], "policies": [],
        "policy_summary": "", "policy_blocked": False,
        "corlo_prompt": "", "prompt_version": "1.0",
        "llm_output": "", "token_estimate": 0, "error": None,
    })

    audit_id  = str(uuid.uuid4())
    tool_info = AI_TOOLS_REGISTRY.get(result["recommended_tool"], {})

    is_blocked     = result.get("policy_blocked", False)
    policy_summary = result.get("policy_summary", "")

    stored_role = (req.role or "").strip() or "general"

    conn = get_db()
    conn.execute(
        """INSERT INTO audit_log
            (id, created_at, raw_input, intent, industry,
             recommended_tool, tool_reason, tool_confidence,
             policy_flags, retrieved_policies, final_prompt,
             prompt_version, model_used, output, token_estimate,
             system_version, policy_blocked, policy_summary, role)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            audit_id, datetime.utcnow().isoformat(),
            result["user_input"], result["intent"], result["industry"],
            result["recommended_tool"], result["tool_reason"],
            result["tool_confidence"], json.dumps(result["policy_flags"]),
            json.dumps(result["policies"]), result["corlo_prompt"],
            result.get("prompt_version", "1.0"),
            result["recommended_tool"], result["llm_output"],
            result["token_estimate"], SYSTEM_VERSION,
            1 if is_blocked else 0,
            policy_summary,
            stored_role,
        ),
    )
    conn.commit()
    conn.close()

    return {
        "audit_id":          audit_id,
        "intent":            result["intent"],
        "industry":          result["industry"],
        "recommended_tool":  result["recommended_tool"],
        "tool_reason":       result["tool_reason"],
        "tool_confidence":   result["tool_confidence"],
        "tool_alternatives":                result["tool_alternatives"],
        "tool_alternative_reasons":         result.get("tool_alternative_reasons", []),
        "tool_alternative_confidence_pcts": result.get("tool_alternative_confidence_pcts", []),
        "tool_confidence_pct":              result.get("tool_confidence_pct", 0),
        "tool_confidence_explanation":      result.get("tool_confidence_explanation", ""),
        "tool_icon":         tool_info.get("icon", "🤖"),
        "tool_category":     tool_info.get("category", ""),
        "tool_url":          tool_info.get("url", ""),
        "tool_is_internal":  tool_info.get("is_internal", False),
        "policy_flags":      result["policy_flags"],
        "policies":          result["policies"],
        "policy_summary":    policy_summary,
        "policy_blocked":    is_blocked,
        "corlo_prompt":      result["corlo_prompt"],
        "prompt_version":    result.get("prompt_version", "1.0"),
        "output":            result["llm_output"],
        "token_estimate":    result["token_estimate"],
        "role":              req.role,
        "task_type":         req.task_type,
        "data_sensitivity":  req.data_sensitivity,
    }


# ══════════════════════════════════════════════════════════════════════════════
# CLARIFICATION  —  assess whether the task description is clear enough,
#                   and if not, return 1-3 short yes/no or one-word questions.
# ══════════════════════════════════════════════════════════════════════════════
@router.post("/api/clarify")
async def check_clarity(req: ClarifyRequest):
    """
    Evaluate whether the user's task description is specific enough for accurate
    tool recommendation. Returns:
      - needs_clarification: bool
      - questions: list[str]  (1-4 short, task-specific questions)
    """
    if not req.user_input.strip():
        raise HTTPException(400, "Input cannot be empty")

    system_msg = (
        "You are a helpful assistant trying to understand what a user wants to do "
        "so you can recommend the right AI tool for them. "
        "Return ONLY valid JSON — no markdown, no extra text."
    )

    user_msg = (
        f"A user described their task as:\n\"{req.user_input}\"\n\n"
        "Analyse the task carefully and ask ONLY the questions that are genuinely needed "
        "to recommend the right tool. Use your judgement:\n"
        "- If the task is simple and clear, ask 1 question or none at all.\n"
        "- If the task is moderately ambiguous, ask 2-3 questions.\n"
        "- Only ask more if there are truly that many distinct unclear parts.\n\n"
        "For each unclear part you find, write one short question asking the user to clarify "
        "that specific part. Only ask about things they actually wrote — never bring up new topics.\n\n"
        "If everything they wrote is already clear, return needs_clarification=false and no questions.\n\n"
        "Return ONLY this JSON:\n"
        "{\"needs_clarification\": true or false, \"questions\": []}"
    )

    try:
        raw = call_llm(system_msg, user_msg, max_tokens=300, temperature=0.2)
        raw = raw.replace("```json", "").replace("```", "").strip()
        data = json.loads(raw)
        needs = bool(data.get("needs_clarification", False))
        questions = [str(q).strip() for q in (data.get("questions") or []) if str(q).strip()]
        return {"needs_clarification": needs, "questions": questions}
    except Exception:
        return {"needs_clarification": False, "questions": []}


@router.post("/api/clarify-merge")
async def merge_clarification(req: ClarifyAnswerRequest):
    """
    Merge the user's original input with Q&A answers into an enriched description
    that the /api/run endpoint can use for accurate tool recommendation.
    """
    if not req.user_input.strip():
        raise HTTPException(400, "Input cannot be empty")

    qa_pairs = "\n".join(
        f"Q: {q}\nA: {a}"
        for q, a in zip(req.questions, req.answers)
        if q.strip() and a.strip()
    )

    system_msg = (
        "You are a task description enricher for an enterprise AI tool recommender. "
        "Combine the original task description with the user's clarification answers "
        "into a single, clear, specific task description. "
        "Output ONLY the enriched description — no preamble, no explanation, no JSON."
    )
    user_msg = (
        f"ORIGINAL TASK DESCRIPTION:\n{req.user_input}\n\n"
        f"CLARIFICATION Q&A:\n{qa_pairs}\n\n"
        "Write a single enriched task description that incorporates all the above details. "
        "Keep it natural and concise (2-4 sentences max)."
    )

    try:
        enriched = call_llm(system_msg, user_msg, max_tokens=300, temperature=0.1)
        return {"enriched_input": enriched.strip()}
    except Exception:
        combined = req.user_input
        if qa_pairs:
            combined += " " + " ".join(req.answers)
        return {"enriched_input": combined.strip()}


# ══════════════════════════════════════════════════════════════════════════════
# REFINEMENT  —  user adds a comment on the existing LLM output
# ══════════════════════════════════════════════════════════════════════════════
@router.post("/api/refine")
async def refine_output(req: RefinementRequest):
    """
    Accept a user comment and revise the CORLO prompt accordingly.
    The revised prompt is returned so the frontend can display it
    and the user can copy it into the recommended AI tool.
    """
    if not req.comment.strip():
        raise HTTPException(400, "Comment cannot be empty")

    system_msg = (
        f"You are an expert prompt engineer helping a {req.role} working on a "
        f"{req.task_type} task in the {req.industry} industry. "
        f"Data sensitivity: {req.data_sensitivity}. "
        f"Target tool: {req.recommended_tool}. "
        "Your job is to revise an existing CORLO prompt based on the user's feedback. "
        "The CORLO prompt is structured in 5 sections: ROLE, CONTEXT, OBJECTIVE, LIMITATIONS, OUTPUT. "
        "Apply the user's comment precisely — change only what they ask, keep everything else. "
        "Return only the complete revised CORLO prompt — no preamble, no explanation."
    )
    user_msg = (
        f"ORIGINAL USER REQUEST:\n{req.user_input}\n\n"
        f"CURRENT CORLO PROMPT (what you are revising):\n{req.corlo_prompt}\n\n"
        f"USER FEEDBACK / REVISION COMMENT:\n{req.comment}\n\n"
        "Instructions:\n"
        "- Read the comment carefully — it tells you exactly what to change in the prompt.\n"
        "- If they say 'make it shorter' → tighten the OBJECTIVE and OUTPUT sections.\n"
        "- If they say 'add X' → insert it in the most appropriate section.\n"
        "- If they say 'focus on Y' → adjust the OBJECTIVE and OUTPUT accordingly.\n"
        "- Keep the 5-section structure (ROLE, CONTEXT, OBJECTIVE, LIMITATIONS, OUTPUT).\n"
        "- Return the complete revised prompt only."
    )

    try:
        revised = call_llm(system_msg, user_msg, max_tokens=2000)
    except Exception as e:
        raise HTTPException(500, f"LLM refinement failed: {str(e)}")

    return {
        "audit_id":       req.audit_id,
        "revised_output": revised,
    }


# ══════════════════════════════════════════════════════════════════════════════
# FEEDBACK
# ══════════════════════════════════════════════════════════════════════════════
@router.post("/api/feedback")
async def submit_feedback(req: FeedbackRequest):
    conn = get_db()
    conn.execute(
        "INSERT INTO feedback VALUES (?,?,?,?,?,?)",
        (str(uuid.uuid4()), req.audit_id, req.rating, req.comment, req.issue_type, datetime.utcnow().isoformat())
    )
    conn.commit()
    conn.close()
    return {"status": "ok"}


# ══════════════════════════════════════════════════════════════════════════════
# POLICIES
# ══════════════════════════════════════════════════════════════════════════════
@router.post("/api/upload-policy")
async def upload_policy(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename or ""
    ext      = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # ── Extract text based on file type ────────────────────────────────────
    if ext == "pdf":
        text = _extract_pdf_text(content)
    elif ext == "docx":
        text = _extract_docx_text(content)
    else:
        # Plain text / markdown fallback
        text = content.decode("utf-8", errors="ignore")

    if not text or not text.strip():
        raise HTTPException(400, "Could not extract any text from the uploaded file. "
                                 "Ensure the file contains readable text.")

    chunks    = [text[i:i+500] for i in range(0, len(text), 500) if text[i:i+500].strip()]
    if not chunks:
        raise HTTPException(400, "File appears empty after text extraction.")

    ids       = [f"{filename}_{i}" for i in range(len(chunks))]
    metadatas = [{"source": filename, "chunk": i} for i in range(len(chunks))]

    try:
        existing = policy_collection.get(where={"source": filename})
        if existing["ids"]:
            policy_collection.delete(ids=existing["ids"])
    except Exception:
        pass

    policy_collection.add(documents=chunks, ids=ids, metadatas=metadatas)
    return {"status": "ok", "filename": filename, "chunks_indexed": len(chunks)}


@router.get("/api/policies")
async def list_policies():
    try:
        result  = policy_collection.get()
        sources = list(set(m.get("source", "unknown") for m in (result.get("metadatas") or [])))
        return {"sources": sources, "total_chunks": len(result.get("ids") or [])}
    except Exception:
        return {"sources": [], "total_chunks": 0}


@router.delete("/api/policies/{filename}")
async def delete_policy(filename: str):
    try:
        existing = policy_collection.get(where={"source": filename})
        if existing["ids"]:
            policy_collection.delete(ids=existing["ids"])
            return {"status": "ok"}
        return {"status": "not_found"}
    except Exception as e:
        raise HTTPException(500, str(e))


# ══════════════════════════════════════════════════════════════════════════════
# TOOLS REGISTRY
# ══════════════════════════════════════════════════════════════════════════════
@router.get("/api/tools")
async def get_tools():
    return {
        name: {
            "description": (info.get("description") or "")[:200],
            "category":    info.get("category", ""),
            "icon":        info.get("icon", ""),
            "best_for":    (info.get("best_for") or [])[:8],
            "url":         info.get("url", ""),
            "is_internal": info.get("is_internal", False),
            "raw_data":    info.get("raw_data", {}),
        }
        for name, info in AI_TOOLS_REGISTRY.items()
    }


# ══════════════════════════════════════════════════════════════════════════════
# AUDIT & ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════
@router.get("/api/audit")
async def get_audit_log(limit: int = 20):
    conn = get_db()
    rows = conn.execute("SELECT * FROM audit_log ORDER BY created_at DESC LIMIT ?", (limit,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@router.patch("/api/audit/{audit_id}")
async def update_audit_log(audit_id: str, payload: dict):
    allowed_fields = {"raw_input", "intent", "industry", "recommended_tool", "final_prompt", "output"}
    updates = {k: v for k, v in payload.items() if k in allowed_fields}
    if not updates:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail="No valid fields to update.")
    conn = get_db()
    set_clause = ", ".join(f"{k} = ?" for k in updates)
    values     = list(updates.values()) + [audit_id]
    conn.execute(f"UPDATE audit_log SET {set_clause} WHERE id = ?", values)
    conn.commit()
    conn.close()
    return {"status": "ok", "audit_id": audit_id}


@router.get("/api/analytics")
async def get_analytics():
    conn = get_db()
    total          = conn.execute("SELECT COUNT(*) as c FROM audit_log").fetchone()["c"]
    intents        = conn.execute("SELECT intent, COUNT(*) as c FROM audit_log GROUP BY intent ORDER BY c DESC").fetchall()
    tools          = conn.execute("SELECT recommended_tool, COUNT(*) as c FROM audit_log GROUP BY recommended_tool ORDER BY c DESC").fetchall()
    industries     = conn.execute("SELECT industry, COUNT(*) as c FROM audit_log GROUP BY industry ORDER BY c DESC LIMIT 5").fetchall()
    avg_rating     = conn.execute("SELECT AVG(rating) as r FROM feedback").fetchone()["r"]
    feedback_count = conn.execute("SELECT COUNT(*) as c FROM feedback").fetchone()["c"]
    issue_types    = conn.execute(
        "SELECT issue_type, COUNT(*) as c FROM feedback WHERE issue_type != '' GROUP BY issue_type ORDER BY c DESC"
    ).fetchall()
    low_rated = conn.execute("""
        SELECT a.intent, a.recommended_tool, f.issue_type, f.comment
        FROM feedback f JOIN audit_log a ON f.audit_id = a.id
        WHERE f.rating <= 2
        ORDER BY f.created_at DESC LIMIT 10
    """).fetchall()
    token_trend = conn.execute(
        "SELECT created_at, token_estimate FROM audit_log ORDER BY created_at DESC LIMIT 10"
    ).fetchall()
    conn.close()

    return {
        "total_runs":     total,
        "avg_rating":     round(avg_rating, 1) if avg_rating else None,
        "feedback_count": feedback_count,
        "intents":        [dict(r) for r in intents],
        "tools":          [dict(r) for r in tools],
        "industries":     [dict(r) for r in industries],
        "issue_types":    [dict(r) for r in issue_types],
        "low_rated_runs": [dict(r) for r in low_rated],
        "token_trend":    [dict(r) for r in token_trend],
    }


# ══════════════════════════════════════════════════════════════════════════════
# ANALYTICS DASHBOARD  (unauthorized, time + role filtered)
# ══════════════════════════════════════════════════════════════════════════════
def _has_column(conn, table: str, column: str) -> bool:
    cols = [r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    return column in cols


def _fill_timeline(rows, since, now, period):
    """Return a continuous list of {label, count} buckets with zero-fill."""
    from datetime import timedelta
    data = dict(rows)
    result = []
    if period == "day":
        current = since.replace(minute=0, second=0, microsecond=0)
        while current <= now:
            bucket  = current.strftime("%H")
            display = current.strftime("%H:00")
            result.append({"label": display, "count": data.get(bucket, 0)})
            current += timedelta(hours=1)
    else:
        current = since.replace(hour=0, minute=0, second=0, microsecond=0)
        while current <= now:
            bucket  = current.strftime("%Y-%m-%d")
            display = current.strftime("%b %d")
            result.append({"label": display, "count": data.get(bucket, 0)})
            current += timedelta(days=1)
    return result


@router.get("/api/analytics-dashboard")
async def get_analytics_dashboard(period: str = "day", role: str = "all"):
    """
    Unauthorized analytics endpoint.
    period: 'day' | 'week' | 'month'
    role:   'all' | any role string (partial LIKE match)
    """
    from datetime import timedelta

    now = datetime.utcnow()
    if period == "week":
        since      = now - timedelta(weeks=1)
        prev_since = since - timedelta(weeks=1)
        tl_fmt     = "%Y-%m-%d"
    elif period == "month":
        since      = now - timedelta(days=30)
        prev_since = since - timedelta(days=30)
        tl_fmt     = "%Y-%m-%d"
    else:  # day (default)
        since      = now - timedelta(days=1)
        prev_since = since - timedelta(days=1)
        tl_fmt     = "%H"

    since_str      = since.isoformat()
    prev_since_str = prev_since.isoformat()

    conn = get_db()

    # Only filter by role if the dedicated role column exists.
    # Never use raw_input as a role fallback — it contains full task descriptions.
    # Build role filter clause safely using parameterized queries
    role_filter_sql  = ""
    role_filter_args = []
    if role != "all" and role.strip():
        role_filter_sql  = " AND LOWER(role) LIKE ?"
        role_filter_args = [f"%{role.lower()}%"]

    base_args      = [since_str]      + role_filter_args
    prev_base_args = [prev_since_str, since_str] + role_filter_args

    # ── Totals ────────────────────────────────────────────────────────────────
    total = conn.execute(
        f"SELECT COUNT(*) as c FROM audit_log WHERE created_at >= ?{role_filter_sql}",
        base_args
    ).fetchone()["c"]

    prev_total = conn.execute(
        f"SELECT COUNT(*) as c FROM audit_log WHERE created_at >= ? AND created_at < ?{role_filter_sql}",
        prev_base_args
    ).fetchone()["c"]

    change_pct = None
    if prev_total and prev_total > 0:
        change_pct = round((total - prev_total) / prev_total * 100)

    # ── By role ───────────────────────────────────────────────────────────────
    by_role_rows = conn.execute(
        "SELECT role, COUNT(*) as count "
        "FROM audit_log "
        "WHERE created_at >= ? "
        "  AND role IS NOT NULL AND TRIM(role) != '' "
        + role_filter_sql +
        " GROUP BY role ORDER BY count DESC LIMIT 15",
        base_args
    ).fetchall()
    by_role = [{"role": r["role"].strip().title() if r["role"].strip().lower() == "general" else r["role"].strip(), "count": r["count"]} for r in by_role_rows]

    # ── By intent ─────────────────────────────────────────────────────────────
    by_intent_rows = conn.execute(
        f"SELECT intent, COUNT(*) as count FROM audit_log "
        f"WHERE created_at >= ? AND (intent IS NOT NULL AND intent != ''){role_filter_sql} "
        f"GROUP BY intent ORDER BY count DESC LIMIT 10",
        base_args
    ).fetchall()
    total_intent = sum(r["count"] for r in by_intent_rows) or 1
    by_intent = [
        {"label": r["intent"] or "—", "count": r["count"],
         "total_pct": round(r["count"] / total_intent * 100)}
        for r in by_intent_rows
    ]

    # ── By tool ───────────────────────────────────────────────────────────────
    by_tool_rows = conn.execute(
        f"SELECT recommended_tool, COUNT(*) as count FROM audit_log "
        f"WHERE created_at >= ? AND (recommended_tool IS NOT NULL AND recommended_tool != ''){role_filter_sql} "
        f"GROUP BY recommended_tool ORDER BY count DESC LIMIT 10",
        base_args
    ).fetchall()
    total_tool = sum(r["count"] for r in by_tool_rows) or 1
    by_tool = [
        {"label": r["recommended_tool"] or "—", "count": r["count"],
         "total_pct": round(r["count"] / total_tool * 100)}
        for r in by_tool_rows
    ]

    # ── Blocked runs ──────────────────────────────────────────────────────────
    blocked = 0
    blocked = conn.execute(
        f"SELECT COUNT(*) as c FROM audit_log "
        f"WHERE created_at >= ? AND policy_blocked = 1{role_filter_sql}",
        base_args
    ).fetchone()["c"]

    # ── Timeline ─────────────────────────────────────────────────────────────
    tl_rows = conn.execute(
        f"SELECT strftime('{tl_fmt}', created_at) as bucket, COUNT(*) as count "
        f"FROM audit_log WHERE created_at >= ?{role_filter_sql} "
        f"GROUP BY bucket ORDER BY bucket ASC",
        base_args
    ).fetchall()
    timeline = _fill_timeline([(r["bucket"], r["count"]) for r in tl_rows], since, now, period)

    conn.close()

    return {
        "period":       period,
        "role_filter":  role,
        "total_runs":   total,
        "change_pct":   change_pct,
        "by_role":      by_role,
        "by_intent":    by_intent,
        "by_tool":      by_tool,
        "blocked_runs": blocked,
        "timeline":     timeline,
    }


# ══════════════════════════════════════════════════════════════════════════════
# PROMPT VERSIONS
# ══════════════════════════════════════════════════════════════════════════════
@router.get("/api/prompt-versions")
async def get_prompt_versions():
    conn = get_db()
    rows = conn.execute(
        "SELECT id, version, intent, industry, change_note, created_at, created_by FROM prompt_versions ORDER BY created_at DESC"
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@router.get("/api/prompt-versions/{version_id}")
async def get_prompt_version(version_id: str):
    conn = get_db()
    row  = conn.execute("SELECT * FROM prompt_versions WHERE id = ?", (version_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "Version not found")
    return dict(row)


@router.post("/api/prompt-versions")
async def create_prompt_version(req: PromptVersionRequest):
    conn = get_db()
    last = conn.execute("SELECT version FROM prompt_versions ORDER BY created_at DESC LIMIT 1").fetchone()
    if last:
        try:
            major, minor = last["version"].split(".")
            new_version  = f"{major}.{int(minor) + 1}"
        except Exception:
            new_version = "1.1"
    else:
        new_version = "1.0"

    vid = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO prompt_versions VALUES (?,?,?,?,?,?,?,?)",
        (vid, new_version, req.intent, req.industry, req.template,
         req.change_note, datetime.utcnow().isoformat(), "user")
    )
    conn.commit()
    conn.close()
    return {"status": "ok", "id": vid, "version": new_version}


@router.post("/api/upload-tools-registry")
async def upload_tools_registry(file: UploadFile = File(...)):
    import service  # import the module, not the dict

    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ("xlsx", "xlsm", "xls"):
        raise HTTPException(400, "Only Excel files (.xlsx, .xlsm, .xls) are supported")

    content = await file.read()

    try:
        service.reload_tools_registry(excel_bytes=content)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Could not read Excel file: {str(e)}")

    if not service.AI_TOOLS_REGISTRY:
        raise HTTPException(400, "File was read but no tools were found. Ensure your sheet has a tool name column and at least one data row.")

    return {"status": "ok", "tools_loaded": len(service.AI_TOOLS_REGISTRY)}



@router.patch("/api/audit/{audit_id}")
async def update_audit_log(audit_id: str, req: AuditUpdateRequest):
    conn = get_db()
    row = conn.execute("SELECT * FROM audit_log WHERE id = ?", (audit_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "Audit log not found")

    current = dict(row)

    new_raw_input    = req.raw_input    if req.raw_input is not None else current.get("raw_input", "")
    new_final_prompt = req.final_prompt if req.final_prompt is not None else current.get("final_prompt", "")
    new_output       = req.output       if req.output is not None else current.get("output", "")

    conn.execute(
        "UPDATE audit_log SET raw_input = ?, final_prompt = ?, output = ? WHERE id = ?",
        (new_raw_input, new_final_prompt, new_output, audit_id),
    )
    conn.commit()

    updated = conn.execute("SELECT * FROM audit_log WHERE id = ?", (audit_id,)).fetchone()
    conn.close()
    return dict(updated)



from pathlib import Path
from fastapi.responses import Response

@router.get("/saml/metadata")
async def saml_metadata():
    cert = Path("saml/sp.crt").read_text()
    cert = cert.replace("-----BEGIN CERTIFICATE-----", "") \
               .replace("-----END CERTIFICATE-----", "") \
               .replace("\n", "").strip()

    base_url = "https://ai-navigator-ashpbzhbcmgeerbt.northeurope-01.azurewebsites.net/"  # ← local for now

    xml = f"""<?xml version="1.0"?>
<md:EntityDescriptor
    xmlns:md="urn:oasis:names:tc:SAML:2.0:metadata"
    xmlns:ds="http://www.w3.org/2000/09/xmldsig#"
    entityID="{base_url}/saml/metadata">
  <md:SPSSODescriptor
      AuthnRequestsSigned="false"
      WantAssertionsSigned="true"
      protocolSupportEnumeration="urn:oasis:names:tc:SAML:2.0:protocol">
    <md:KeyDescriptor use="signing">
      <ds:KeyInfo>
        <ds:X509Data>
          <ds:X509Certificate>{cert}</ds:X509Certificate>
        </ds:X509Data>
      </ds:KeyInfo>
    </md:KeyDescriptor>
    <md:NameIDFormat>urn:oasis:names:tc:SAML:1.1:nameid-format:emailAddress</md:NameIDFormat>
    <md:AssertionConsumerService
        Binding="urn:oasis:names:tc:SAML:2.0:bindings:HTTP-POST"
        Location="{base_url}/saml/acs"
        index="1"/>
  </md:SPSSODescriptor>
</md:EntityDescriptor>"""

    return Response(content=xml, media_type="application/xml")